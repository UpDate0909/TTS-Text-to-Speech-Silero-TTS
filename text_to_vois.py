# text_to_vois.py
# Полный скрипт для озвучки текста через Silero TTS (v5_ru) с GUI
# Работает с .txt, .docx, .doc → вывод в MP3 рядом с исходником

import os
import re
import sys
import shutil
import logging
import subprocess
import urllib.request
from pathlib import Path
from typing import List
from datetime import datetime

# === Импорт tkinter для GUI ===
try:
    import tkinter as tk
    from tkinter import filedialog, ttk, messagebox, scrolledtext
except ImportError:
    print("tkinter не найден. Убедитесь, что вы используете стандартный Python.")
    sys.exit(1)

# === Настройка логирования СРАЗУ ===
def setup_logging():
    """Создаем лог-файл в папке со скриптом"""
    if getattr(sys, 'frozen', False):
        # Если запущен как .exe
        log_dir = Path(sys.executable).parent
    else:
        # Если запущен как .py
        log_dir = Path(__file__).parent

    log_file = log_dir / "text_to_vois.log"

    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s | %(levelname)s | %(message)s',
        handlers=[
            logging.FileHandler(log_file, encoding='utf-8', mode='w'),
            logging.StreamHandler(sys.stdout)
        ]
    )
    return logging.getLogger()

logger = setup_logging()
logger.info("=" * 60)
logger.info("ЗАПУСК ПРОГРАММЫ")
logger.info("=" * 60)

# === Поиск ffmpeg ===
def find_ffmpeg():
    """Ищем ffmpeg в PATH или стандартных местах"""
    # Проверяем PATH
    ffmpeg_cmd = "ffmpeg.exe" if sys.platform == "win32" else "ffmpeg"
    if shutil.which(ffmpeg_cmd):
        logger.info(f"✅ ffmpeg найден в PATH: {shutil.which(ffmpeg_cmd)}")
        return shutil.which(ffmpeg_cmd)

    # Проверяем стандартные места Windows
    common_paths = [
        r"C:\ffmpeg\bin\ffmpeg.exe",
        r"C:\Program Files\ffmpeg\bin\ffmpeg.exe",
        Path.home() / "ffmpeg" / "bin" / "ffmpeg.exe",
    ]

    for path in common_paths:
        if Path(path).exists():
            logger.info(f"✅ ffmpeg найден: {path}")
            return str(path)

    return None

FFMPEG_PATH = find_ffmpeg()

if not FFMPEG_PATH:
    logger.critical("❌ ffmpeg.exe не найден!")
    logger.critical("Скачайте: https://www.gyan.dev/ffmpeg/builds/")
    logger.critical("Распакуйте и добавьте в PATH или положите в C:\\ffmpeg\\bin\\")
    print("\n" + "=" * 60)
    print("⚠️  ОШИБКА: ffmpeg не найден!")
    print("=" * 60)
    print("\nИнструкция:")
    print("1. Скачайте: https://www.gyan.dev/ffmpeg/builds/")
    print("2. Распакуйте архив")
    print("3. Скопируйте папку bin в C:\\ffmpeg\\bin\\")
    print("   (должен быть файл C:\\ffmpeg\\bin\\ffmpeg.exe)")
    print("\nИли добавьте ffmpeg в PATH системы")
    print("=" * 60)
    input("\nНажмите Enter для выхода...")
    sys.exit(1)

# Проверим работоспособность
try:
    result = subprocess.run(
        [FFMPEG_PATH, "-version"],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        timeout=5,
        creationflags=subprocess.CREATE_NO_WINDOW if sys.platform == "win32" else 0
    )
    if result.returncode != 0:
        raise Exception(f"ffmpeg вернул код {result.returncode}")
    logger.info("✅ ffmpeg работает корректно")
except Exception as e:
    logger.critical(f"❌ ffmpeg не запускается: {e}")
    print("\n⚠️  ffmpeg найден, но не работает!")
    print(f"Ошибка: {e}")
    input("\nНажмите Enter для выхода...")
    sys.exit(1)

# === Импорты с проверкой ===
logger.info("Проверка зависимостей...")
missing_packages = []

try:
    import torch
    logger.info("✅ torch")
except ImportError:
    missing_packages.append("torch")
    logger.error("❌ torch не установлен")

# Проверяем pydub отдельно
try:
    from pydub import AudioSegment
    # Устанавливаем путь к ffmpeg ДО использования AudioSegment
    AudioSegment.converter = FFMPEG_PATH
    logger.info("✅ pydub (импорт и настройка ffmpeg успешны)")
except ImportError as e:
    if "pydub" in str(e):
        missing_packages.append("pydub")
        logger.error("❌ pydub не установлен")
    else:
        # Ошибка при импорте другого модуля внутри pydub
        logger.error(f"❌ pydub импортирован, но ошибка при настройке: {e}")
        # Если ошибка связана с audioop, pydub установлен, но не может работать без ffmpeg
        # Но мы уже установили ffmpeg, значит, это внутренняя проблема pydub
        logger.info("✅ pydub (импорт успешен, возможна ошибка при использовании из-за отсутствия audioop)")
except ModuleNotFoundError as e:
    if "audioop" in str(e):
        logger.error(f"❌ pydub импортирован, но не может найти 'audioop': {e}")
        logger.error("   Это известная проблема в Python 3.13. Убедитесь, что ffmpeg установлен и указан в AudioSegment.converter.")
        logger.info("✅ pydub (импорт успешен, но возможна ошибка при использовании из-за отсутствия audioop)")
    else:
        missing_packages.append("pydub")
        logger.error(f"❌ pydub импортирован, но ошибка: {e}")
except Exception as e:
    logger.error(f"❌ pydub импортирован, но ошибка при настройке: {e}")
    # Если ошибка связана с audioop, pydub всё равно может работать с ffmpeg
    if "audioop" not in str(e):
        missing_packages.append("pydub")

try:
    from tqdm import tqdm
    logger.info("✅ tqdm")
except ImportError:
    missing_packages.append("tqdm")
    logger.error("❌ tqdm не установлен")

try:
    import soundfile as sf
    logger.info("✅ soundfile")
except ImportError:
    missing_packages.append("soundfile")
    logger.error("❌ soundfile не установлен")

try:
    from docx import Document
    logger.info("✅ python-docx")
except ImportError:
    missing_packages.append("python-docx")
    logger.error("❌ python-docx не установлен")

try:
    import win32com.client
    logger.info("✅ pywin32")
except ImportError:
    missing_packages.append("pywin32")
    logger.error("❌ pywin32 не установлен")

if missing_packages:
    logger.critical(f"❌ Отсутствуют библиотеки: {', '.join(missing_packages)}")
    print("\n" + "=" * 60)
    print("⚠️  ОШИБКА: Не установлены зависимости!")
    print("=" * 60)
    print("\nВыполните в командной строке:")
    print("pip install torch torchaudio pydub soundfile python-docx pywin32 tqdm scipy")
    print("=" * 60)
    input("\nНажмите Enter для выхода...")
    sys.exit(1)

logger.info("✅ Все зависимости установлены")

# === Настройка pydub ===
# Уже выполнено выше: AudioSegment.converter = FFMPEG_PATH

# === Голоса ===
SPEAKERS_INFO = {
    'aidar':    {'gender': 'мужской', 'style': 'спокойный, чёткий'},
    'baya':     {'gender': 'женский', 'style': 'энергичный, яркий'},  # Изменено: теперь женский
    'eugene':   {'gender': 'мужской', 'style': 'глубокий, бархатистый'},
    'kseniya':  {'gender': 'женский', 'style': 'деловой, уверенный'},
    'xenia':    {'gender': 'женский', 'style': 'мягкий, дружелюбный'},
}

# === Чтение файлов ===
def read_txt(file_path: Path) -> str:
    """Читаем .txt с автоопределением кодировки"""
    encodings = ['utf-8', 'windows-1251', 'cp1251']
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                text = f.read()
                logger.info(f"Файл прочитан с кодировкой: {enc}")
                return text
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Не удалось прочитать файл ни с одной из кодировок: {encodings}")

def read_docx(file_path: Path) -> str:
    """Читаем .docx"""
    try:
        doc = Document(file_path)
        text = '\n'.join([para.text for para in doc.paragraphs])
        logger.info(f"Прочитано параграфов: {len(doc.paragraphs)}")
        return text
    except Exception as e:
        logger.error(f"Ошибка чтения .docx: {e}")
        raise

def read_doc(file_path: Path) -> str:
    """Читаем .doc через Word COM"""
    logger.info("Открываем Word для чтения .doc...")
    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(file_path.absolute()))
        text = doc.Content.Text
        doc.Close(False)
        word.Quit()
        logger.info("Файл .doc прочитан успешно")
        return text
    except Exception as e:
        logger.error(f"Ошибка чтения .doc: {e}")
        if word:
            try:
                word.Quit()
            except:
                pass
        raise ValueError(f"Не удалось открыть .doc: {e}\nУбедитесь, что установлен Microsoft Word.")

def read_text_file(file_path: Path) -> str:
    """Универсальный читатель"""
    ext = file_path.suffix.lower()
    logger.info(f"Чтение файла: {file_path.name} ({ext})")

    if ext == '.txt':
        return read_txt(file_path)
    elif ext == '.docx':
        return read_docx(file_path)
    elif ext == '.doc':
        return read_doc(file_path)
    else:
        raise ValueError(f"Неподдерживаемый формат: {ext}")

# === Разбивка на предложения ===
def split_into_sentences(text: str) -> List[str]:
    """Разбиваем текст на предложения по . ! ? — – - : ... , (многоточие, запятая)"""
    # Убираем лишние пробелы и переводы строк, заменяя на один пробел
    text = re.sub(r'\s+', ' ', text).strip()

    # Заменяем все \s+, которые идут после последовательности [.!?—–\-:,]+ или ... на специальный маркер.
    # Это позволяет корректно обработать последовательности знаков.
    marker = '\n---SPLIT---\n'
    # Паттерн: ([.!?—–\-:,]+|\.{3})\s+
    # Заменим на: \1 + marker
    processed_text = re.sub(r'([.!?—–\-:,]+|\.{3})\s+', r'\1' + marker, text)

    # Теперь разбиваем по маркеру
    sentences = processed_text.split(marker)

    # Убираем пустые строки и лишние пробелы
    result = [s.strip() for s in sentences if s.strip()]
    logger.info(f"Текст разбит на {len(result)} предложений")
    return result

def group_sentences(sentences: List[str], max_chars: int = 4900) -> List[str]:
    """Группируем предложения в чанки до max_chars символов"""
    chunks = []
    current = []
    length = 0

    for s in sentences:
        s_len = len(s)
        if length + s_len + 1 <= max_chars:
            current.append(s)
            length += s_len + 1
        else:
            if current:
                chunks.append(' '.join(current))
            current = [s]
            length = s_len

    if current:
        chunks.append(' '.join(current))

    logger.info(f"Создано {len(chunks)} чанков для озвучки")
    return chunks

# === Загрузка модели Silero ===
def get_silero_model_path():
    """Получаем путь к модели (скачиваем при необходимости)"""
    model_dir = Path.home() / ".cache" / "silero"
    model_dir.mkdir(parents=True, exist_ok=True)
    model_path = model_dir / "v5_ru.pt"

    if not model_path.exists():
        logger.info("Модель не найдена, начинаем загрузку...")
        logger.info("Размер: ~100 МБ, это может занять несколько минут")
        try:
            def show_progress(block_num, block_size, total_size):
                downloaded = block_num * block_size
                if total_size > 0:
                    percent = min(100, downloaded * 100 / total_size)
                    print(f"\rЗагрузка модели: {percent:.1f}%", end='')

            url = 'https://models.silero.ai/models/tts/ru/v5_ru.pt'
            urllib.request.urlretrieve(url, str(model_path), show_progress)
            print()  # Новая строка после прогресса
            logger.info(f"✅ Модель загружена: {model_path}")
        except Exception as e:
            logger.error(f"Ошибка загрузки модели: {e}")
            raise
    else:
        logger.info(f"Модель найдена: {model_path}")

    return model_path

def load_silero_model():
    """Загружаем модель Silero TTS"""
    try:
        model_path = get_silero_model_path()
        logger.info("Инициализация модели...")
        model = torch.package.PackageImporter(str(model_path)).load_pickle("tts_models", "model")
        model.to(torch.device('cpu'))
        logger.info("✅ Модель готова к работе")
        return model
    except Exception as e:
        logger.error(f"Ошибка загрузки модели: {e}")
        raise

# === Генерация аудио ===
def generate_audio_chunk(model, text: str, speaker: str, sample_rate: int, output_path: Path):
    """Генерируем аудио для одного чанка"""
    try:
        audio = model.apply_tts(
            text=text,
            speaker=speaker,
            sample_rate=sample_rate,
            put_accent=True,
            put_yo=True
        )
        sf.write(str(output_path), audio.numpy(), sample_rate)
    except Exception as e:
        logger.error(f"Ошибка генерации аудио: {e}")
        raise

# === Склейка в MP3 (через ffmpeg напрямую) ===
def convert_and_concatenate(wav_files: List[Path], output_mp3: Path):
    """Склеиваем WAV и конвертируем в MP3 через ffmpeg напрямую"""
    import subprocess
    import tempfile

    logger.info("Начинаем склейку и конвертацию через ffmpeg...")
    try:
        # Создаём временный .txt файл с путями к WAV
        with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.txt', encoding='utf-8') as f:
            for wav in wav_files:
                # В ffmpeg списке файлов нужно экранировать пути
                f.write(f"file '{wav.absolute()}'\n")
            list_file = f.name

        # Команда ffmpeg для склейки и конвертации
        cmd = [
            FFMPEG_PATH,  # Используем путь, найденный в начале
            "-f", "concat",
            "-safe", "0",
            "-i", list_file,
            "-c:a", "libmp3lame",  # для MP3
            "-b:a", "192k",        # битрейт
            "-y",                  # перезаписать
            str(output_mp3)
        ]

        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        if result.returncode != 0:
            raise Exception(f"ffmpeg вернул ошибку: {result.stderr}")

        logger.info(f"✅ Аудио сохранено: {output_mp3}")
        # Получим длительность через ffmpeg
        cmd_duration = [FFMPEG_PATH, "-i", str(output_mp3), "-f", "null", "-"]
        result_dur = subprocess.run(cmd_duration, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        duration_line = [l for l in result_dur.stderr.split('\n') if 'Duration' in l]
        if duration_line:
            duration = duration_line[0].split(',')[0].replace('Duration: ', '').strip()
            logger.info(f"Длительность: {duration}")

    except Exception as e:
        logger.error(f"Ошибка конвертации: {e}")
        raise
    finally:
        # Удаляем временный .txt файл
        try:
            os.unlink(list_file)
        except:
            pass

# === Основной класс GUI ===
class TTSApp:
    def __init__(self, root):
        self.root = root
        self.root.title("TTS (Text-to-Speech) на базе Silero TTS\nМодифицированная версия от UpDate0909")  # Изменено: заголовок
        self.root.geometry("700x600")
        self.root.resizable(True, True)

        # === Переменные ===
        self.file_path = tk.StringVar()
        self.selected_voice = tk.StringVar(value="xenia — женский, мягкий, дружелюбный")  # Изменено: теперь полное имя

        # === Интерфейс ===
        self.create_widgets()

    def get_selected_speaker_key(self):
        selected = self.selected_voice.get()
        key = selected.split(" — ")[0]
        return key

    def create_widgets(self):
        # === Заголовок ===
        title = tk.Label(self.root, text="TTS (Text-to-Speech) на базе Silero TTS\nМодифицированная версия от UpDate0909", font=("Arial", 14, "bold"))
        title.pack(pady=10)

        # === Путь к файлу ===
        file_frame = tk.Frame(self.root)
        file_frame.pack(fill="x", padx=20, pady=5)

        tk.Label(file_frame, text="Файл (.txt, .docx, .doc):").pack(anchor="w")
        file_entry = tk.Entry(file_frame, textvariable=self.file_path, width=60)
        file_entry.pack(side="left", fill="x", expand=True, padx=(0, 10))
        tk.Button(file_frame, text="Обзор", command=self.browse_file).pack(side="right")

        # === Выбор голоса ===
        voice_frame = tk.Frame(self.root)
        voice_frame.pack(fill="x", padx=20, pady=5)

        tk.Label(voice_frame, text="Голос:").pack(anchor="w")
        voice_selector = ttk.Combobox(voice_frame, textvariable=self.selected_voice, state="readonly", width=50)
        voice_selector['values'] = [f"{k} — {v['gender']}, {v['style']}" for k, v in SPEAKERS_INFO.items()]
        voice_selector.current(4)  # по умолчанию xenia
        voice_selector.pack(side="left", fill="x", expand=True, padx=(0, 10))

        # === Кнопка запуска ===
        tk.Button(self.root, text="Начать озвучку", command=self.start_processing, bg="lightgreen", font=("Arial", 12)).pack(pady=20)

        # === Прогресс-бар ===
        self.progress = ttk.Progressbar(self.root, orient="horizontal", length=600, mode="determinate")
        self.progress.pack(pady=5)

        # === Логи ===
        log_frame = tk.Frame(self.root)
        log_frame.pack(fill="both", expand=True, padx=20, pady=10)

        tk.Label(log_frame, text="Логи:").pack(anchor="w")
        self.log_text = scrolledtext.ScrolledText(log_frame, height=15, state="normal")
        self.log_text.pack(fill="both", expand=True)

        # === Статус ===
        self.status = tk.Label(self.root, text="Готов", bd=1, relief="sunken", anchor="w")
        self.status.pack(side="bottom", fill="x")

    def browse_file(self):
        # Изменено: по умолчанию "Все файлы"
        filetypes = (
            ("Все файлы", "*.*"),
            ("Текстовые файлы", "*.txt"),
            ("Документы Word", "*.docx"),
            ("Старые документы Word", "*.doc"),
        )
        filename = filedialog.askopenfilename(title="Выберите файл", filetypes=filetypes)
        if filename:
            self.file_path.set(filename)

    def log(self, message):
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()
        logger.info(message)

    def start_processing(self):
        # === Проверки ===
        file_path_str = self.file_path.get()
        if not file_path_str:
            messagebox.showerror("Ошибка", "Пожалуйста, выберите файл.")
            return

        file = Path(file_path_str)
        if not file.exists():
            messagebox.showerror("Ошибка", f"Файл не найден: {file}")
            return

        if file.suffix.lower() not in ['.txt', '.docx', '.doc']:
            messagebox.showerror("Ошибка", f"Неподдерживаемый формат: {file.suffix}")
            return

        # Проверка ffmpeg
        if FFMPEG_PATH != shutil.which("ffmpeg"):
            if not Path(FFMPEG_PATH).exists():
                messagebox.showerror("Ошибка", f"ffmpeg не найден по пути: {FFMPEG_PATH}")
                return
        else:
            self.log(f"✅ ffmpeg найден в PATH: {FFMPEG_PATH}")

        # === Запуск процесса ===
        self.status.config(text="Запуск...")
        self.progress['value'] = 0

        try:
            self.process_file(file)
        except Exception as e:
            logger.error(f"КРИТИЧЕСКАЯ ОШИБКА: {e}")
            messagebox.showerror("Ошибка", f"Произошла ошибка: {e}")

    def process_file(self, file: Path):
        import torch

        self.log(f"Чтение файла: {file.name}")
        text = read_text_file(file)

        if not text.strip():
            raise ValueError("Файл пуст или не содержит текста!")

        self.log(f"Прочитано символов: {len(text)}")
        sentences = split_into_sentences(text)
        # Изменено: теперь чанки — это по 1 предложению
        chunks = [s for s in sentences if s]  # Просто каждое предложение — отдельный чанк
        self.log(f"Текст разбит на {len(chunks)} чанков (по одному предложению)")

        # === Загрузка модели ===
        self.log("Загрузка модели Silero...")
        model = load_silero_model()
        sample_rate = 48000
        self.log("✅ Модель загружена")

        # === Временная папка ===
        temp_dir = file.parent / "temp_tts_chunks"
        temp_dir.mkdir(exist_ok=True)
        wav_files = []

        # === Генерация аудио ===
        # Считаем только те чанки, которые содержат текст
        valid_chunks = [c for c in chunks if re.search(r'[а-яА-ЯёЁa-zA-Z0-9]', c)]
        total = len(valid_chunks)
        self.log(f"Найдено {total} валидных чанков для озвучки (остальные пропущены).")

        for i, chunk in enumerate(valid_chunks):
            # ЛОГИРУЕМ текст предложения перед обработкой
            logger.info(f"Обрабатывается предложение: {chunk}")
            wav_path = temp_dir / f"audio_{i+1:03d}.wav"
            speaker_key = self.get_selected_speaker_key()  # ✅ Извлекаем ключ голоса
            generate_audio_chunk(model, chunk, speaker_key, sample_rate, wav_path)
            wav_files.append(wav_path)

            # Обновляем прогресс
            self.progress['value'] = (i + 1) / total * 100
            self.status.config(text=f"Озвучка... {i+1}/{total}")
            self.log(f"Обработан чанк {i+1}/{total}")
            self.root.update_idletasks()

        self.log("✅ Озвучка завершена")

        # === Конвертация в MP3 ===
        self.status.config(text="Конвертация в MP3...")
        # Изменено: формат имени файла
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        speaker_name = self.get_selected_speaker_key()
        output_mp3 = file.parent / f"{file.stem}_{speaker_name}_{timestamp}.mp3"  # Формат: имя_файла_голос_дата_время.mp3
        convert_and_concatenate(wav_files, output_mp3)
        self.log(f"✅ Аудио сохранено: {output_mp3.name}")

        # === Удаление временных файлов ===
        shutil.rmtree(temp_dir, ignore_errors=True)
        self.log("✅ Временные файлы удалены")

        self.status.config(text="Готово!")
        self.progress['value'] = 100
        messagebox.showinfo("Успех", f"Аудиофайл успешно создан:\n{output_mp3.name}")

# === Основной процесс (для обратной совместимости, если запускается как консольный скрипт) ===
def main():
    print("\n" + "=" * 60)
    print("  🎙️  SILERO TTS: Преобразование текста в речь")
    print("=" * 60 + "\n")
    print("Запуск GUI версии...")
    root = tk.Tk()
    app = TTSApp(root)
    root.mainloop()

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n❌ Прервано пользователем")
        logger.info("Программа прервана пользователем")
    except Exception as e:
        logger.exception("КРИТИЧЕСКАЯ ОШИБКА:")
        print(f"\n❌ КРИТИЧЕСКАЯ ОШИБКА: {e}")
        print("\nПодробности в файле text_to_vois.log")
        input("\nНажмите Enter для выхода...")